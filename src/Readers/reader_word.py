from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import docx
import os
from docx.document import Document as _Document
from src.model.paragraph import Paragraph as ParagraphHexa
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

class WordReader:

    def __init__(self, path):
        self.path = path
        self.paragraphs = self.get_paragraphs()
    def iter_block_items(self, parent):
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Unsupported parent type")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def get_paragraphs(self):
        if not os.path.exists(self.path):
            raise FileNotFoundError(f"The file {self.path} does not exist.")
        try:
            doc = docx.Document(self.path)
            paragraph_objects = []
            paragraph_id = 0
            page_id = 1  # Example page ID
            total_characters = 0
            for block in self.iter_block_items(doc):
                if isinstance(block, Paragraph):
                    paragraph_info = self.extract_paragraph_info(block)
                    if paragraph_info:  # Only append if paragraph is not empty
                        page_id = self.estimate_page_number(total_characters)
                        p_obj = ParagraphHexa(text=paragraph_info['text'], font_style=paragraph_info['style'], id_=paragraph_id, page_id=page_id)
                        #print(f"Found paragraph: {paragraph_info['style']}...")  # DEBUG
                        paragraph_objects.append(p_obj)
                        paragraph_id += 1
                        total_characters += len(paragraph_info['text'])
                elif isinstance(block, Table):
                    table_paragraph, table_style = self.table_to_paragraph(block)
                    if table_paragraph.strip():  # Check if table paragraph is not empty
                        #print(f"Found table. Predominant style: {table_style}")  # DEBUG
                        p_obj = ParagraphHexa(text=table_paragraph, font_style=table_style, id_=paragraph_id, page_id=page_id)
                        paragraph_objects.append(p_obj)
                        paragraph_id += 1
            return paragraph_objects
        except Exception as e:
            raise ValueError(f"Error reading the .docx file. Original error: {str(e)}")

        
    def determine_predominant_style(self, styles):
        # Count the occurrences of each style
        style_counts = {}
        for style in styles:
            if style in style_counts:
                style_counts[style] += 1
            else:
                style_counts[style] = 1

        # Find the style with the highest count
        predominant_style = max(style_counts, key=style_counts.get, default="None")
        return predominant_style

    def estimate_page_number(self, total_characters):
        avg_chars_per_page = 2000  
        return total_characters // avg_chars_per_page + 1

    def extract_paragraph_info(self, paragraph):
        # Check if paragraph is empty
        if not paragraph.text.strip():
            return None  # Return None for empty paragraphs

        paragraph_style = paragraph.style.name if paragraph.style else 'None'

        runs = []
        for run in paragraph.runs:
            run_details = {
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size.pt if run.font.size else None,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline
            }
            runs.append(run_details)

        return {
            'text': paragraph.text,
            'style': paragraph_style,
            'runs': runs
        }



    def table_to_paragraph(self, table):
        table_text = ""
        table_styles = set()

        for row in table.rows:
            for cell in row.cells:
                cell_text = ""
                for paragraph in cell.paragraphs:
                    paragraph_style = paragraph.style.name if paragraph.style else 'None'
                    table_styles.add(paragraph_style)

                    for run in paragraph.runs:
                        cell_text += run.text

                    cell_text += " "
                table_text += cell_text.strip() + " | "  # Add a separator for cells
            table_text = table_text.strip() + "\n"  # Add a newline for rows

        predominant_style = self.determine_predominant_style(table_styles)

        return table_text.strip(), predominant_style

    def print_paragraphs_and_tables(self):
        try:
            print("start")
            doc_items = self.get_paragraphs()
            for item in doc_items:
                if 'paragraph' in item:
                    print("Paragraph:", item['paragraph']['text'])
                elif 'table' in item:
                    print("Table:")
                    for row in item['table']:
                        for cell in row:
                            for paragraph in cell:
                                print("   Cell Paragraph:", paragraph['text'])
                print('-' * 40)  # separator for clarity

        except Exception as e:
            print(f"Error: {str(e)}")



# path = "/Users/quent1/Documents/Hexamind/ILLUMIO/Illumio3011/test_data/Illumio_Core_REST_API_Developer_Guide_23.3.docx" # Replace with the path to your Word document
# if not os.path.exists(path):
#     print("The file does not exist. Please check the path.")
# else:
#     word_reader = WordReader(path)
#     word_reader.get_paragraphs()



################# LEGACY CODE ################




# from __future__ import (
#     absolute_import, division, print_function, unicode_literals
# )
# import docx
# import os
# from src.model.paragraph import Paragraph as ParagraphHexa
# from docx import Document
# from docx.document import Document as _Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.table import _Cell, Table
# from docx.text.paragraph import Paragraph

# class WordReader:

#     def __init__(self, path):
#         self.path = path
#         #self.paragraphs = self.get_paragraphs()

#     def iter_block_items(self,parent):
#         """
#         Generate a reference to each paragraph and table child within *parent*,
#         in document order. Each returned value is an instance of either Table or
#         Paragraph. *parent* would most commonly be a reference to a main
#         Document object, but also works for a _Cell object, which itself can
#         contain paragraphs and tables.
#         """
#         if isinstance(parent, _Document):
#             parent_elm = parent.element.body
#             # print(parent_elm.xml)
#         elif isinstance(parent, _Cell):
#             parent_elm = parent._tc
#         else:
#             raise ValueError("something's not right")

#         for child in parent_elm.iterchildren():
#             if isinstance(child, CT_P):
#                 yield Paragraph(child, parent)
#             elif isinstance(child, CT_Tbl):
#                 yield Table(child, parent)        
            
#     def get_paragraphs(self):
#         """
#         Fetches paragraphs from a Word document.

#         Returns:
#             list: List of Paragraph objects from the document.
#         """
#         if not os.path.exists(self.path):
#             raise FileNotFoundError(f"The file {self.path} does not exist.")

#         try:
#             doc = docx.Document(self.path)
#             doc_paragraphs = []
            
#             for block in self.iter_block_items(doc):
#                 if isinstance(block, Paragraph):
#                     paragraph_text = block.text
#                     doc_paragraphs.append(paragraph_text)
#                 elif isinstance(block, Table):
#                     paragraph_table = self.table_to_paragraph(block)
#                     doc_paragraphs.append(paragraph_table)
                    
#             return self.to_paragraph_objects(doc_paragraphs)
#         except Exception as e:
#             raise ValueError(f"Error reading the .docx file. Original error: {str(e)}")
        
        
#     def to_paragraph_objects(self, doc_paragraphs):
#         """
#         Convert docx paragraphs to Paragraph objects for further processing.
#         """
#         paragraph_objects = []
#         for idx, paragraph in enumerate(doc_paragraphs):
#             style = self.determine_style(paragraph)
#             if paragraph:  # Skip empty paragraphs
#                 print(idx, paragraph)
#                 p_obj = ParagraphHexa( font_style="normal", id_=idx, page_id=1, text=paragraph)
#                 paragraph_objects.append(p_obj)

#         paragraphs = self.rearrange_paragraphs(paragraph_objects)
#         return paragraphs
            
            
            
#     def get_table_content(self, table):
#         """
#         Returns the content of a table as a nested list.
#         """
#         table_content = []
#         for row in table.rows:
#             row_content = []
#             for cell in row.cells:
#                 row_content.append(cell.text)
#             table_content.append(row_content)
#         return table_content


#     def determine_style(self, paragraph):
#         """
#         Determines the style of the paragraph based on its attributes.

#         Returns:
#             str: Style of the paragraph.
#         """
#         # Check for heading styles first
#         if paragraph.style.name.startswith('Heading 1'):
#             return "title1"
#         elif paragraph.style.name.startswith('Heading 2'):
#             return "title2"
#         elif paragraph.style.name.startswith('Heading 3'):
#             return "title3"
#         elif paragraph.style.name.startswith('Heading 4'):
#             return "title4"
#         elif paragraph.style.name.startswith('Heading 5'):
#             return "title5"
    
#         # If not a heading, check the runs within the paragraph
#         for run in paragraph.runs:
#             font = run.font
#             fontname = font.name
#             size = font.size
        
#             # Convert size to points (from twips)
#             if size:
#                 size_in_points = size.pt

#                 # Map based on font name and size as in the PDF reader
#                 if fontname == "XFQKGD+Consolas":
#                     return "code"
#                 elif (size_in_points >= 9 and size_in_points < 11.5) or fontname == "Wingdings-Regular":
#                     return "content"    
#         # If none of the above conditions match, default to 'content'
#         return "content"
    


#     def table_to_paragraph(self, block):
#         table = block
#         paragraph = []
#         for row in table.rows:
#             for cell in row.cells:
#                 for p in cell.paragraphs:
#                     paragraph.append(p.text + ' ') 
#         return paragraph
        
#     def get_table_paragraphs(self, table):
#         """
#         Returns the table content as a list of Paragraph objects.
#         """
#         table_paragraphs = []
#         counter = 1  # Initialize the counter
#         for row in table.rows:
#             for cell in row.cells:
#                 paragraph = Paragraph(text=cell.text, font_style="normal", id_=counter, page_id=0)
#                 table_paragraphs.append(paragraph)
#         return table_paragraphs
        

#     def rearrange_paragraphs(self, paragraphs : [Paragraph]):
#         #associate paragraphs with the same font style
#         i = 0
#         while i < len(paragraphs):
#             paragraphs[i] = paragraphs[i].rearrange_paragraph()
#             i+=1
#         return paragraphs

#     def print_paragraphs_and_tables(self):
#         """
#         Print paragraphs and tables from the Word document.
#         """
#         try:
#             doc_paragraphs = self.get_paragraphs()
#             for paragraph in doc_paragraphs:
#                 print("Paragraph:", paragraph.text)

#         except Exception as e:
#             print(f"Error: {str(e)}")


#     def display_paragraphs(self):
#         """
#         Prints the paragraphs from the document to the console.
#         """
#         for paragraph in self.paragraphs:
#             print(paragraph.text)
#             print('-' * 40)  # separator for clarity
            

#     def display_tables(self):
#         """
#         Prints the tables from the document to the console.
#         """
#         for table in self.tables:
#             table_content = self.get_table_content(table)
#             for row in table_content:
#                 print(row)
#             print('-' * 40)  # separator for clarity
            
            

# if __name__ == "__main__":
#     path = "/Users/quent1/Documents/Hexamind/ILLUMIO/Illumio3011/test_data/Illumio_Core_REST_API_Developer_Guide_23.3.docx"  # Replace with the path to your Word document
#     if not os.path.exists(path):
#         print("The file does not exist. Please check the path.")
#     else:
#         word_reader = WordReader(path)
#         word_reader.print_paragraphs_and_tables()